from tkinter import *
from tkinter.font import BOLD
from tkinter import messagebox
import tkinter.ttk as tkrttk
from docx.shared import Pt
from docx2pdf import convert
#from tkcalendar import Calendar, DateEntry
from datetime import date
from threading import Timer
import pymongo
from pymongo import MongoClient
import docx as docx
import os

from docx2pdf import convert
cluster=MongoClient("mongodb+srv://overseaseducation943:asdfgf;lkjhj@cluster0.ssarbxh.mongodb.net/?retryWrites=true&w=majority")
db=cluster.overseaseducation
collection=db.details

def onClear():
    s_college_name.delete(0, END)
    s_enrollment_number.delete(0, END)
    s_name.delete(0, END)
    s_branch.set('');
    s_student_mobile.delete(0, END)
    s_parent_mobile.delete(0, END)
    drop.set('');
    s_fee.delete(0, END)
    s_discount.delete(0, END)
    s_discount.insert(0,"0")
    s_fee.insert(0,"0")
    T.delete('1.0','end')
    T2.delete('1.0','end')
    email_entry.delete(0,END)

def onSubmit():
    try:
        a=int(s_fee.get())
        b=int(s_discount.get())
        final_fee_value=a-(b)
        if(s_enrollment_number.get()==''):
            messagebox.showinfo("Invalid!","Enter the Enrollment Number")
        else:
            
            post={'enrollment_number':s_enrollment_number.get(),'name':s_name.get(),'clg_name':s_college_name.get(),
                  ' branch':s_branch.get(),
                  'stuphno':s_student_mobile.get(),'parphno':s_parent_mobile.get(),
    'course':str(course1.get())
    ,'fee':s_fee.get()
    ,'discount':s_discount.get()
    ,'reason':T.get('1.0','end-1c'),
    #'date':cal.get(),
    'address':T2.get('1.0','end-1c'),
    'fee_payer':s_namep.get(),
    'email':email_entry.get()}
            #collection.insert_one(post)
            doc = docx.Document()
            style = doc.styles['Normal']
            font = style.font
            font.size = Pt(16)
            doc.add_heading("Phoenix Overseas Education Services",0)
            doc.add_heading("Office Copy",2)
            
            doc.add_paragraph("Date(YYYY:MM:DD):\t\t"+str(date.today()))
            doc.add_paragraph("Enrollment number:\t"+s_enrollment_number.get())
            doc.add_paragraph("Name:\t\t"+s_name.get())
            doc.add_paragraph("Email:\t\t"+email_entry.get())
            doc.add_paragraph("College Name:\t\t"+s_college_name.get())
            doc.add_paragraph("Branch:\t\t"+s_branch.get())
            doc.add_paragraph("Student Mobile Number:"+s_student_mobile.get())
            doc.add_paragraph("Parent Mobile Number:"+s_parent_mobile.get())
            doc.add_paragraph("Course:\t\t"+str(course1.get()))
            doc.add_paragraph("Fee:\t\t"+s_fee.get())
            doc.add_paragraph("Discount:\t"+s_discount.get())
            doc.add_paragraph("Refered by:\t "+T.get('1.0','end-1c'))
            doc.add_paragraph("Address:\t\t"+T2.get('1.0','end-1c'))
            ddd=str(final_fee_value)
            #print(a,b,final_fee_value)
            doc.add_paragraph('Final Fees:\t\t'+ddd)
            doc.add_paragraph("Fee Received by:\t"+s_namep.get())
            doc.add_paragraph("\n\n\t\t\t\t\tSignature of the receiver")
            doc.add_page_break()
            doc.add_heading("Phoenix Overseas Education Services",0)
            doc.add_heading("Student Copy",2)
            
            doc.add_paragraph("Date(YYYY:MM:DD):\t\t"+str(date.today()))
            doc.add_paragraph("Enrollment number:\t"+s_enrollment_number.get())
            doc.add_paragraph("Name:\t\t"+s_name.get())
            doc.add_paragraph("Email:\t\t"+email_entry.get())
            doc.add_paragraph("College Name:\t\t"+s_college_name.get())
            doc.add_paragraph("Branch:\t\t"+s_branch.get())
            doc.add_paragraph("Student Mobile Number:"+s_student_mobile.get())
            doc.add_paragraph("Parent Mobile Number:"+s_parent_mobile.get())
            doc.add_paragraph("Course:\t\t"+str(course1.get()))
            doc.add_paragraph("Fee:\t\t"+s_fee.get())
            doc.add_paragraph("Discount:\t"+s_discount.get())
            doc.add_paragraph("Refered by:\t "+T.get('1.0','end-1c'))
            doc.add_paragraph("Address:\t\t"+T2.get('1.0','end-1c'))
            ddd=str(final_fee_value)
            #print(a,b,final_fee_value)
            doc.add_paragraph('Final Fees:\t\t'+ddd)
            doc.add_paragraph("Fee Received by:\t"+s_namep.get())
            doc.add_paragraph("\n\n\t\t\t\t\tSignature of the receiver")
            
            doc.save("C:\\"+'Fee-Receipts\\Docx\\'+str(s_enrollment_number.get())+"out.docx")
            
            convert("C:\\"+'Fee-Receipts\\Docx\\'+str(s_enrollment_number.get())+"out.docx","C:\\"+'Fee-Receipts\\pdf\\'+str(s_enrollment_number.get())+"out.pdf")
            messagebox.showinfo("Success","The receipt was successfully generated")
            #doca.save("E:\\"+str(s_enrollment_number.get())+"out.pdf")
           # os.startfile((str(s_enrollment_number.get())+'out.pdf'), "print")
    except ValueError as v:
            messagebox.showinfo("Error","Invalid Input")

        

root=Tk()
root.geometry("900x800")
root.title("Student-Pay Form")

#For enrollment number
enrollment_number=Label(root, text='Enrollment Number:',font=('bold',10))
enrollment_number.place(x=20,y=30)
s_enrollment_number=Entry()
s_enrollment_number.place(x=220,y=30)


#For Name
name=Label(root, text='Name:',font=('bold',10))
name.place(x=20,y=60)
s_name=Entry()
s_name.place(x=220,y=60)


#For College
college_name=Label(root, text='College Name:',font=('bold',10))
college_name.place(x=20,y=90)
s_college_name=Entry()
s_college_name.place(x=220,y=90)


#For branch
branch=Label(root, text='Branch:',font=('bold',10))
branch.place(x=20,y=120)
branch_name=StringVar()
s_branch=tkrttk.Combobox(root,textvariable=branch_name)
s_branch["values"]=["CSE","ECE","EEE","CIVIL","MECH","IT"]
s_branch.place(x=220,y=120)

#For Mobile
student_mobile=Label(root, text='Student Mobile Number:',font=('bold',10))
student_mobile.place(x=20,y=150)
s_student_mobile=Entry()
s_student_mobile.place(x=220,y=150)

parent_mobile=Label(root, text='Parent Mobile Number:',font=('bold',10))
parent_mobile.place(x=20,y=180)
s_parent_mobile=Entry()
s_parent_mobile.place(x=220,y=180)

#For Course as an drop down
course=Label(root, text='Course:',font=('bold',10))
course.place(x=20,y=210)
course1=StringVar()
drop=tkrttk.Combobox(root,textvariable=course1)
drop["values"]=["F1","F2","F3","F4","F5","F6","F7","F8","F9","F10","F11"]
drop.place(x=220,y=210)

#For fees 
fee=Label(root, text='Fees:',font=('bold',10))
fee.place(x=20,y=240)
s_fee=Entry()
s_fee.insert(0,"0")
s_fee.place(x=220,y=240)

#For discount if any
discount=Label(root, text='Discount:',font=('bold',10))
discount.place(x=20,y=270)
s_discount=Entry()
s_discount.insert(0,"0")
s_discount.place(x=220,y=270)

#Reason
reason=Label(root, text='Reason:',font=('bold',10))
reason.place(x=20,y=300)
T = Text(root, height = 5, width = 30)
T.place(x=220,y=300)

#Date
#date=Label(root, text='Date:',font=('bold',10))
#date.place(x=20,y=410)
#cal = DateEntry(root, width= 16, background= "magenta3", foreground= "white",bd=2)
#cal.place(x=220,y=410)


#Address
address=Label(root, text='Address:',font=('bold',10))
address.place(x=20,y=450)
T2 = Text(root, height = 4, width = 30)
T2.place(x=220,y=450)


#For Fee recievedby
name=Label(root, text='Fee received by:',font=('bold',10))
name.place(x=20,y=550)
s_namep=Entry()
s_namep.place(x=220,y=550)

# email
email_s=Label(root,text="Enter your mail : ",font=('bold',10))
email_s.place(x=20,y=600)
email_entry=Entry()
email_entry.place(x=220,y=600)


#Submit Button
sub=Button(root,text='Submit',command=onSubmit)
sub.place(x=20,y=680)

#Clear Button
sub=Button(root,text='Clear',command=onClear)
sub.place(x=120,y=680)
root.mainloop()




